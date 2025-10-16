from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Optional
from datetime import datetime
import os


class ExportService:
    """Service for exporting question papers to PDF and Word formats."""
    
    def __init__(self, output_dir: str = "./exports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export_to_pdf(
        self,
        questions: List[Dict],
        paper_config: Dict,
        include_answers: bool = False
    ) -> str:
        """Export question paper to PDF format."""
        
        filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_JUSTIFY
        )
        
        # Header
        institution_name = paper_config.get('institution_name', 'Institution Name')
        elements.append(Paragraph(institution_name, title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Paper title
        paper_title = paper_config.get('title', 'Question Paper')
        elements.append(Paragraph(paper_title, heading_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Paper info table
        info_data = [
            ['Time:', f"{paper_config.get('duration_minutes', 60)} minutes", 
             'Total Marks:', str(paper_config.get('total_marks', 100))],
        ]
        
        info_table = Table(info_data, colWidths=[1*inch, 2*inch, 1*inch, 1*inch])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Instructions
        if paper_config.get('instructions'):
            elements.append(Paragraph("<b>Instructions:</b>", heading_style))
            elements.append(Paragraph(paper_config['instructions'], question_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Horizontal line
        elements.append(Spacer(1, 0.1*inch))
        
        # Questions
        for idx, question in enumerate(questions, 1):
            # Question number and text
            q_text = f"<b>Q{idx}.</b> {question['question_text']}"
            if question.get('suggested_marks'):
                q_text += f" <i>({question['suggested_marks']} marks)</i>"
            
            elements.append(Paragraph(q_text, question_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Options for MCQ
            if question['question_type'] in ['mcq', 'true_false']:
                options = []
                for opt_key in ['option_a', 'option_b', 'option_c', 'option_d']:
                    if question.get(opt_key):
                        opt_label = opt_key.split('_')[1].upper()
                        options.append(f"({opt_label}) {question[opt_key]}")
                
                for option in options:
                    elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{option}", question_style))
                
                elements.append(Spacer(1, 0.1*inch))
            
            # Answer space for written questions
            elif question['question_type'] in ['short_answer', 'fill_blank']:
                elements.append(Spacer(1, 0.3*inch))
            elif question['question_type'] in ['long_answer', 'programming']:
                elements.append(Spacer(1, 0.5*inch))
            
            elements.append(Spacer(1, 0.15*inch))
        
        # Answer key (if requested)
        if include_answers:
            elements.append(PageBreak())
            elements.append(Paragraph("<b>Answer Key</b>", title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            for idx, question in enumerate(questions, 1):
                answer_text = f"<b>Q{idx}.</b> {question['correct_answer']}"
                elements.append(Paragraph(answer_text, question_style))
                
                if question.get('explanation'):
                    exp_text = f"<i>Explanation: {question['explanation']}</i>"
                    elements.append(Paragraph(exp_text, question_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(elements)
        
        return filepath
    
    def export_to_word(
        self,
        questions: List[Dict],
        paper_config: Dict,
        include_answers: bool = False
    ) -> str:
        """Export question paper to Word format."""
        
        filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header
        institution_name = paper_config.get('institution_name', 'Institution Name')
        heading = doc.add_heading(institution_name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Paper title
        paper_title = paper_config.get('title', 'Question Paper')
        title = doc.add_heading(paper_title, level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Paper info
        info_para = doc.add_paragraph()
        info_para.add_run(f"Time: {paper_config.get('duration_minutes', 60)} minutes").bold = True
        info_para.add_run(f"\t\tTotal Marks: {paper_config.get('total_marks', 100)}").bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Instructions
        if paper_config.get('instructions'):
            doc.add_heading('Instructions:', level=3)
            doc.add_paragraph(paper_config['instructions'])
            doc.add_paragraph()
        
        # Horizontal line
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Questions
        for idx, question in enumerate(questions, 1):
            # Question text
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{idx}. {question['question_text']}")
            q_run.font.size = Pt(11)
            
            if question.get('suggested_marks'):
                marks_run = q_para.add_run(f" ({question['suggested_marks']} marks)")
                marks_run.italic = True
                marks_run.font.size = Pt(10)
            
            # Options for MCQ
            if question['question_type'] in ['mcq', 'true_false']:
                for opt_key in ['option_a', 'option_b', 'option_c', 'option_d']:
                    if question.get(opt_key):
                        opt_label = opt_key.split('_')[1].upper()
                        opt_para = doc.add_paragraph(f"({opt_label}) {question[opt_key]}")
                        opt_para.paragraph_format.left_indent = Inches(0.5)
            
            # Answer space
            elif question['question_type'] in ['short_answer', 'fill_blank']:
                doc.add_paragraph()
                doc.add_paragraph('_' * 60)
            elif question['question_type'] in ['long_answer', 'programming']:
                for _ in range(5):
                    doc.add_paragraph('_' * 80)
            
            doc.add_paragraph()  # Spacing between questions
        
        # Answer key
        if include_answers:
            doc.add_page_break()
            doc.add_heading('Answer Key', level=1)
            
            for idx, question in enumerate(questions, 1):
                ans_para = doc.add_paragraph()
                ans_run = ans_para.add_run(f"Q{idx}. ")
                ans_run.bold = True
                ans_para.add_run(question['correct_answer'])
                
                if question.get('explanation'):
                    exp_para = doc.add_paragraph()
                    exp_run = exp_para.add_run(f"Explanation: {question['explanation']}")
                    exp_run.italic = True
                    exp_para.paragraph_format.left_indent = Inches(0.5)
                
                doc.add_paragraph()
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def export_answer_key_only(self, questions: List[Dict], paper_config: Dict) -> str:
        """Export only the answer key as a separate PDF."""
        
        filename = f"answer_key_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        elements.append(Paragraph("Answer Key", title_style))
        elements.append(Paragraph(paper_config.get('title', 'Question Paper'), styles['Heading2']))
        elements.append(Spacer(1, 0.3*inch))
        
        # Answers
        for idx, question in enumerate(questions, 1):
            q_text = f"<b>Q{idx}.</b> {question['correct_answer']}"
            elements.append(Paragraph(q_text, styles['Normal']))
            
            if question.get('explanation'):
                exp_text = f"<i>{question['explanation']}</i>"
                elements.append(Paragraph(exp_text, styles['Normal']))
            
            elements.append(Spacer(1, 0.15*inch))
        
        doc.build(elements)
        return filepath
